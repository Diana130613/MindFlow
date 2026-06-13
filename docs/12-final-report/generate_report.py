"""
Генератор пояснительной записки для курсового проекта MindFlow.
Запуск: python generate_report.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─── Страница ────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(1.5)
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)

# ─── Стили ───────────────────────────────────────────────────────────────────
styles = doc.styles

def set_style(style_name, font_name='Times New Roman', size=14, bold=False,
              italic=False, align=WD_ALIGN_PARAGRAPH.LEFT,
              space_before=0, space_after=6, first_line=0):
    try:
        st = styles[style_name]
    except KeyError:
        st = styles.add_style(style_name, 1)
    pf = st.paragraph_format
    pf.alignment       = align
    pf.space_before    = Pt(space_before)
    pf.space_after     = Pt(space_after)
    if first_line:
        pf.first_line_indent = Cm(first_line)
    rf = st.font
    rf.name  = font_name
    rf.size  = Pt(size)
    rf.bold  = bold
    rf.italic = italic
    return st

normal_st  = set_style('Normal',      size=14, space_after=6,  first_line=1.25)
h1_st      = set_style('Heading 1',   size=16, bold=True,  align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=6)
h2_st      = set_style('Heading 2',   size=14, bold=True,  space_before=10, space_after=4)
h3_st      = set_style('Heading 3',   size=14, bold=False, italic=True, space_before=6, space_after=4)
title_st   = set_style('Title',       size=18, bold=True,  align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=12)
caption_st = set_style('Caption',     size=12, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=6)
code_st    = set_style('Code',        font_name='Courier New', size=11, space_before=4, space_after=4)

# ─── Хелперы ─────────────────────────────────────────────────────────────────

def add_para(text='', style='Normal', align=None, bold=False, indent=True):
    p = doc.add_paragraph(style=style)
    if align:
        p.alignment = align
    if not indent:
        p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    run.bold = bold
    return p

def add_h1(text):
    p = doc.add_paragraph(style='Heading 1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'
    return p

def add_h2(text):
    p = doc.add_paragraph(style='Heading 2')
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    return p

def add_h3(text):
    p = doc.add_paragraph(style='Heading 3')
    run = p.add_run(text)
    run.bold = False
    run.italic = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    return p

def add_text(text, indent=True):
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(1.25) if indent else Cm(0)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.25 + level * 0.63)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    return p

def add_code(text):
    p = doc.add_paragraph(style='Code')
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    return p

def add_caption(text):
    p = doc.add_paragraph(style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.italic = True
    return p

def add_table_caption(text):
    p = doc.add_paragraph(style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.bold = True
    return p

def add_simple_table(headers, rows, caption=None):
    if caption:
        add_table_caption(caption)
    col_count = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=col_count)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D9D9D9')
        shading.set(qn('w:val'), 'clear')
        c._tc.get_or_add_tcPr().append(shading)
    # Data rows
    for ri, row in enumerate(rows):
        tr = t.rows[ri + 1]
        for ci, cell_val in enumerate(row):
            c = tr.cells[ci]
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = c.paragraphs[0]
            run = p.add_run(str(cell_val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
    doc.add_paragraph()

def page_break():
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# ТИТУЛЬНЫЙ ЛИСТ
# ═══════════════════════════════════════════════════════════════════════════════

def title_run(p, text, size=14, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    p.alignment = align
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.bold = bold

p = doc.add_paragraph()
title_run(p, 'МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ', 12)

p = doc.add_paragraph()
title_run(p, 'Федеральное государственное автономное образовательное учреждение высшего образования', 12)

p = doc.add_paragraph()
title_run(p, '«СЕВЕРО-КАВКАЗСКИЙ ФЕДЕРАЛЬНЫЙ УНИВЕРСИТЕТ»', 13, bold=True)

doc.add_paragraph()

p = doc.add_paragraph()
title_run(p, 'Институт информационных технологий и телекоммуникаций', 12)
p = doc.add_paragraph()
title_run(p, 'Направление подготовки 09.03.04 «Программная инженерия»', 12)
p = doc.add_paragraph()
title_run(p, 'Профиль: «Разработка и сопровождение программного обеспечения»', 12)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
title_run(p, 'ПОЯСНИТЕЛЬНАЯ ЗАПИСКА', 18, bold=True)
p = doc.add_paragraph()
title_run(p, 'к курсовому проекту по дисциплине', 14)
p = doc.add_paragraph()
title_run(p, '«Программная инженерия»', 14, bold=True)

doc.add_paragraph()

p = doc.add_paragraph()
title_run(p, 'Тема: «MindFlow — мобильное приложение для ментального здоровья»', 15, bold=True)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.first_line_indent = Cm(0)
r = p.add_run('Выполнила: студентка группы ПИЖ-б-о-23-2\nХатуаева Дайана Алиевна\n\n')
r.font.name = 'Times New Roman'; r.font.size = Pt(13)
r2 = p.add_run('Руководитель: ___________________\n')
r2.font.name = 'Times New Roman'; r2.font.size = Pt(13)
r3 = p.add_run('Дата защиты: «___» __________ 2026 г.\n')
r3.font.name = 'Times New Roman'; r3.font.size = Pt(13)
r4 = p.add_run('Оценка: ____________________')
r4.font.name = 'Times New Roman'; r4.font.size = Pt(13)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
title_run(p, 'Ставрополь 2026', 14)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# АННОТАЦИЯ
# ═══════════════════════════════════════════════════════════════════════════════
add_h1('АННОТАЦИЯ')
add_text(
    'В данной пояснительной записке описывается курсовой проект «MindFlow» — '
    'мобильное приложение для поддержки ментального здоровья, разработанное в '
    'рамках дисциплины «Программная инженерия» по траектории В (мобильная '
    'разработка). Проект реализован как полноценная клиент-серверная система: '
    'Android-приложение на Kotlin с Jetpack Compose и серверная часть на Java '
    'Spring Boot.'
)
add_text(
    'Документ охватывает все этапы жизненного цикла ПО: от бизнес-анализа '
    'и проектирования требований до реализации, тестирования и развёртывания. '
    'Архитектура системы построена на паттерне PCMEF (Presentation-Control-'
    'Mediator-Entity-Foundation), обеспечивающем строгое разделение '
    'ответственности между компонентами.'
)
add_text(
    'Объём разработки: 12 экранов Android-приложения, 15 REST API эндпоинтов, '
    '289 модульных тестов с покрытием 43,6% (Android) и 58,1% (Backend), '
    'полный комплект проектной документации в 13 разделах.'
)
page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# СОДЕРЖАНИЕ (ручное, т.к. python-docx не генерирует TOC автоматически)
# ═══════════════════════════════════════════════════════════════════════════════
add_h1('СОДЕРЖАНИЕ')

toc_items = [
    ('ВВЕДЕНИЕ', '5'),
    ('1. АНАЛИТИЧЕСКАЯ ЧАСТЬ', '6'),
    ('   1.1. Описание предметной области', '6'),
    ('   1.2. Анализ бизнес-процессов', '7'),
    ('   1.3. SWOT-анализ', '7'),
    ('   1.4. Анализ аналогов', '8'),
    ('   1.5. Customer Journey Map', '8'),
    ('   1.6. Экономическое обоснование (ROI)', '9'),
    ('2. ПРОЕКТНАЯ ЧАСТЬ', '10'),
    ('   2.1. Модель требований', '10'),
    ('   2.2. Domain Model', '11'),
    ('   2.3. Архитектурное проектирование', '12'),
    ('   2.4. Проектирование базы данных', '15'),
    ('   2.5. Детальное проектирование', '17'),
    ('3. РЕАЛИЗАЦИОННАЯ ЧАСТЬ', '19'),
    ('   3.1. Структура проекта', '19'),
    ('   3.2. Реализация серверной части (Backend)', '20'),
    ('   3.3. Реализация Android-клиента', '24'),
    ('   3.4. Рефакторинг и паттерны качества', '27'),
    ('   3.5. Безопасность системы', '29'),
    ('   3.6. REST API', '30'),
    ('4. ТЕСТИРОВАНИЕ', '31'),
    ('   4.1. Стратегия тестирования', '31'),
    ('   4.2. Модульное тестирование бэкенда', '32'),
    ('   4.3. Модульное тестирование Android', '33'),
    ('   4.4. Ручное тестирование', '34'),
    ('5. РАЗВЁРТЫВАНИЕ', '35'),
    ('6. УПРАВЛЕНИЕ ПРОЕКТОМ', '36'),
    ('   6.1. WBS', '36'),
    ('   6.2. Диаграмма Ганта', '37'),
    ('   6.3. Оценка трудозатрат (COCOMO)', '37'),
    ('ЗАКЛЮЧЕНИЕ', '38'),
    ('СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ', '39'),
]

for item, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(item)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(13)
    tab_stop = OxmlElement('w:tab')
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:leader'), 'dot')
    tab.set(qn('w:pos'), '9072')
    tabs.append(tab)
    pPr.append(tabs)
    r2 = p.add_run('\t' + page)
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(13)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# ВВЕДЕНИЕ
# ═══════════════════════════════════════════════════════════════════════════════
add_h1('ВВЕДЕНИЕ')

add_text(
    'Проблемы ментального здоровья приобретают всё большую значимость в современном '
    'обществе. По данным Всемирной организации здравоохранения, депрессия занимает '
    'первое место среди причин инвалидности в мире, а тревожные расстройства '
    'диагностируются у каждого восьмого человека на планете. В России, по данным '
    'Министерства здравоохранения, около 40% населения испытывают симптомы '
    'хронического стресса. При этом доступность квалифицированной психологической '
    'помощи остаётся недостаточной: дефицит специалистов, высокая стоимость сеансов '
    'и социальная стигматизация создают серьёзные барьеры для обращения за помощью.'
)
add_text(
    'Мобильные приложения для ментального здоровья открывают новые возможности '
    'для преодоления этих барьеров. Они обеспечивают доступ к инструментам '
    'самопомощи в любое время и в любом месте, не требуют записи к специалисту '
    'и могут использоваться анонимно. Медитация и практики осознанности, '
    'являющиеся основой таких приложений, имеют доказанную научную эффективность '
    'в снижении тревожности, улучшении качества сна и повышении общего '
    'психоэмоционального благополучия.'
)
add_text(
    'Актуальность разработки приложения MindFlow обусловлена тремя факторами: '
    'растущим спросом на цифровые инструменты ментального здоровья, недостаточным '
    'предложением русскоязычных решений на рынке и образовательной ценностью проекта '
    'как демонстрации полного цикла разработки мобильного клиент-серверного '
    'приложения.'
)

add_h2('Цель и задачи')
add_text(
    'Цель проекта: разработать мобильное приложение для Android, обеспечивающее '
    'пользователям доступ к библиотеке медитаций и инструментам мониторинга '
    'эмоционального состояния.'
)
add_text('Задачи:')
add_bullet('провести анализ предметной области и конкурентной среды;')
add_bullet('спроектировать архитектуру системы на основе паттерна PCMEF;')
add_bullet('реализовать серверную часть на Java Spring Boot с REST API;')
add_bullet('разработать Android-клиент на Kotlin с Jetpack Compose;')
add_bullet('обеспечить офлайн-доступ через локальное кэширование (Room);')
add_bullet('реализовать аутентификацию через JWT-токены;')
add_bullet('покрыть код модульными тестами (>40% по строкам);')
add_bullet('подготовить полный комплект проектной документации.')

add_h2('Объект и предмет исследования')
add_text(
    'Объект исследования: процесс поддержки ментального здоровья с применением '
    'цифровых технологий.'
)
add_text(
    'Предмет исследования: архитектурные паттерны и технологии разработки '
    'мобильных клиент-серверных приложений.'
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# ГЛАВА 1. АНАЛИТИЧЕСКАЯ ЧАСТЬ
# ═══════════════════════════════════════════════════════════════════════════════
add_h1('1. АНАЛИТИЧЕСКАЯ ЧАСТЬ')

add_h2('1.1. Описание предметной области')
add_text(
    'Предметная область проекта — поддержка ментального здоровья через практики '
    'медитации и осознанности (mindfulness). Медитация — это группа техник '
    'тренировки внимания, направленных на достижение состояния повышенной '
    'осознанности и сосредоточенности. Научные исследования подтверждают, что '
    'регулярная практика медитации снижает уровень кортизола (гормона стресса), '
    'улучшает качество сна, повышает концентрацию внимания и развивает '
    'эмоциональный интеллект.'
)
add_text(
    'Ключевые категории контента в приложении:'
)
add_bullet('Медитации осознанности (mindfulness) — наблюдение за дыханием и мыслями без оценки')
add_bullet('Дыхательные практики — структурированные техники дыхания (4-7-8, капалабхати)')
add_bullet('Сканирование тела — систематическое расслабление через внимание к ощущениям')
add_bullet('Визуализации — направляемые образные путешествия для снижения тревоги')
add_bullet('Йога-нидра — медитация на границе сна и бодрствования')

add_text(
    'Дневник настроения — дополнительный инструмент, позволяющий пользователю '
    'фиксировать эмоциональное состояние по шкале от 1 до 10 с текстовыми '
    'заметками. Регулярный мониторинг помогает выявлять паттерны и триггеры '
    'эмоционального состояния.'
)

add_h2('1.2. Анализ бизнес-процессов')
add_text(
    'Основной бизнес-процесс описан посредством IDEF0-диаграммы уровня A-0. '
    'Система получает на вход запросы пользователей на медитацию и данные '
    'о настроении, трансформирует их в персонализированный опыт практики, '
    'управляемый правилами безопасности и принципами UX-дизайна, и выдаёт '
    'на выходе: завершённые сессии медитации, записи настроения и аналитику.'
)

add_simple_table(
    ['Вход', 'Управление', 'Механизм', 'Выход'],
    [
        ['Запрос на медитацию', 'PCMEF-архитектура', 'Android + Spring Boot', 'Сессия медитации'],
        ['Запись настроения', 'JWT-безопасность', 'PostgreSQL + Room', 'Аналитика'],
        ['Данные профиля', 'Material Design 3', 'Retrofit + OkHttp', 'Персонализация'],
    ],
    'Таблица 1 — Контекстная диаграмма IDEF0 A-0'
)

add_h2('1.3. SWOT-анализ')

add_simple_table(
    ['', 'Положительные факторы', 'Отрицательные факторы'],
    [
        ['Внутренние', 'Офлайн-режим, PCMEF-архитектура, JWT-безопасность, Material Design 3', 'Нет аудиофайлов в репозитории, нет push-уведомлений'],
        ['Внешние', 'Рост рынка wellness-приложений (+18% CAGR), дефицит русскоязычных аналогов', 'Сильные конкуренты (Calm, Headspace), монетизация требует инвестиций'],
    ],
    'Таблица 2 — SWOT-анализ проекта MindFlow'
)

add_h2('1.4. Анализ аналогов и конкурентов')

add_simple_table(
    ['Критерий', 'MindFlow', 'Calm', 'Headspace', 'Simple Habit'],
    [
        ['Язык интерфейса', '🇷🇺 Русский', '🇬🇧 Английский', '🇬🇧 Английский', '🇬🇧 Английский'],
        ['Офлайн-режим', '✅ Да', '✅ Да (Premium)', '✅ Да (Premium)', '❌ Нет'],
        ['Открытый код', '✅ GitHub', '❌ Нет', '❌ Нет', '❌ Нет'],
        ['Дневник настроения', '✅ Да', '✅ Да', '✅ Да', '❌ Нет'],
        ['Бесплатный базовый план', '✅ Freemium', '✅ Ограниченно', '✅ Ограниченно', '✅ Ограниченно'],
        ['REST API', '✅ Открытый', '❌ Закрытый', '❌ Закрытый', '❌ Закрытый'],
        ['Цена Premium', '490 руб./мес', '~1 200 руб./мес', '~1 000 руб./мес', '~900 руб./мес'],
    ],
    'Таблица 3 — Сравнительный анализ конкурентов'
)

add_text(
    'Ключевые конкурентные преимущества MindFlow: полностью русскоязычный интерфейс, '
    'открытый код, офлайн-режим в бесплатной версии и более доступная цена Premium-подписки.'
)

add_h2('1.5. Customer Journey Map')
add_text(
    'Для проектирования пользовательского опыта была разработана карта пути '
    'пользователя (CJM). Целевая персона: Алина, 22 года, студентка, испытывающая '
    'стресс перед сессией.'
)

add_simple_table(
    ['Этап', 'Действие', 'Эмоция', 'Боль', 'Решение в MindFlow'],
    [
        ['Осознание', 'Слышит о приложении', '😕 Скептицизм', 'Страх «это не для меня»', 'Экран Welcome с мотивацией'],
        ['Регистрация', 'Создаёт аккаунт', '😐 Нейтральна', 'Сложная форма', '3 поля: имя, email, пароль'],
        ['Первая медитация', '5 мин дыхание 4-7-8', '😌 Спокойствие', 'Боится длинных практик', 'Сортировка по длительности'],
        ['Дневник', 'Записывает оценку 4', '🤔 Интерес', 'Непонятна шкала', 'Emoji + подписи к числам'],
        ['Аналитика', 'Видит прогресс за неделю', '😊 Радость', 'Нет мотивации продолжать', 'График настроения, тренд'],
        ['Рекомендация', 'Советует подруге', '😃 Энтузиазм', '—', '—'],
    ],
    'Таблица 4 — Customer Journey Map (персона Алина)'
)

add_h2('1.6. Экономическое обоснование (ROI)')
add_text(
    'Экономическая оценка проводилась в контексте потенциального коммерческого '
    'запуска приложения на российском рынке wellness-приложений.'
)
add_simple_table(
    ['Показатель', 'Значение'],
    [
        ['Объём рынка (РФ, 2025)', '~4,2 млрд руб.'],
        ['CAGR рынка', '~18%'],
        ['Фактические затраты на разработку', '~93 600 руб.'],
        ['Прогнозируемый ROI (3 года)', '~401%'],
        ['Точка безубыточности', '~3 050 MAU'],
        ['Период выхода на безубыточность', '6–9 месяцев'],
    ],
    'Таблица 5 — Ключевые экономические показатели'
)
add_text(
    'Модель монетизации — Freemium: базовый функционал бесплатно, Premium-подписка '
    'за 490 руб./месяц. При консервативном прогнозе в 5 000 MAU и конверсии 3% '
    'доход первого года составит ~523 500 руб. при затратах ~734 100 руб., '
    'что обеспечивает выход на окупаемость в начале второго года.'
)
page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# ГЛАВА 2. ПРОЕКТНАЯ ЧАСТЬ
# ═══════════════════════════════════════════════════════════════════════════════
add_h1('2. ПРОЕКТНАЯ ЧАСТЬ')

add_h2('2.1. Модель требований к ПО')

add_h3('2.1.1. Акторы системы')
add_simple_table(
    ['Актор', 'Описание', 'Права доступа'],
    [
        ['Гость', 'Неаутентифицированный пользователь', 'Регистрация, вход'],
        ['Пользователь (ROLE_USER)', 'Аутентифицированный пользователь', 'Медитации, дневник настроения, профиль'],
        ['Администратор (ROLE_ADMIN)', 'Привилегированный пользователь', 'Всё + управление контентом'],
    ],
    'Таблица 6 — Акторы системы'
)

add_h3('2.1.2. Основные прецеденты')
add_simple_table(
    ['Код', 'Прецедент', 'Актор', 'Описание'],
    [
        ['UC-01', 'Регистрация', 'Гость', 'Создание нового аккаунта (email + пароль + имя)'],
        ['UC-02', 'Аутентификация', 'Гость', 'Вход с получением JWT-токенов'],
        ['UC-03', 'Просмотр медитаций', 'Пользователь', 'Каталог с фильтром и поиском'],
        ['UC-04', 'Проведение медитации', 'Пользователь', 'Таймер, прогресс-бар, завершение'],
        ['UC-05', 'Запись настроения', 'Пользователь', 'Шкала 1–10, заметка, сохранение'],
        ['UC-06', 'Просмотр аналитики', 'Пользователь', 'График динамики настроения за период'],
        ['UC-07', 'Управление профилем', 'Пользователь', 'Просмотр статистики, выход'],
        ['UC-08', 'Работа в офлайн-режиме', 'Пользователь', 'Доступ к кэшированному контенту'],
    ],
    'Таблица 7 — Прецеденты системы'
)

add_h3('2.1.3. Спецификация ключевого прецедента UC-04')
add_text(
    'Прецедент: «Проведение медитации»'
)
add_bullet('Предусловие: пользователь аутентифицирован, медитация выбрана')
add_bullet('Основной поток: открыть Detail → нажать «Начать» → таймер запущен → прогресс-бар → «Завершить» → сессия сохранена')
add_bullet('Альтернативный поток: нажать «Стоп» → сессия сохраняется как незавершённая')
add_bullet('Постусловие: MeditationSession создана в БД, прогресс пользователя обновлён')
add_bullet('Исключения: отсутствие сети → данные сохраняются локально с флагом sync_pending')

add_h3('2.1.4. Нефункциональные требования')
add_simple_table(
    ['Категория', 'Требование', 'Целевое значение'],
    [
        ['Производительность', 'Время отклика API (LAN)', '≤ 500 мс'],
        ['Производительность', 'Загрузка из Room-кэша', '≤ 100 мс'],
        ['Надёжность', 'Покрытие тестами', '>40% (LINE)'],
        ['Безопасность', 'Хеширование паролей', 'BCrypt, strength=12'],
        ['Совместимость', 'Минимальная версия Android', 'API 26 (Android 8.0)'],
        ['Удобство', 'Соответствие дизайн-системе', 'Material Design 3'],
    ],
    'Таблица 8 — Нефункциональные требования'
)

add_h2('2.2. Domain Model (модель предметной области)')
add_text(
    'Концептуальная модель системы включает 7 бизнес-сущностей с чётко '
    'определёнными атрибутами, бизнес-методами и связями. Модель разработана '
    'с применением техники Domain-Driven Design и отражает язык предметной '
    'области, а не детали реализации.'
)

add_simple_table(
    ['Сущность', 'Ключевые атрибуты', 'Бизнес-методы', 'Связи'],
    [
        ['User', 'id, email, passwordHash, name, role, createdAt', 'isAdmin(), updateProfile()', '→ MeditationSession, MoodEntry, UserProgress'],
        ['Meditation', 'id, title, description, durationMinutes, type, difficulty', 'isForBeginners(), getShortDescription()', '→ MeditationSession, Category'],
        ['MeditationSession', 'id, startedAt, completedAt, durationSeconds, completed', 'markCompleted(), getDurationMinutes()', '→ User, Meditation'],
        ['MoodEntry', 'id, score (1–10), note, recordedAt', 'isPositive(), getMoodLabel(), validate()', '→ User'],
        ['Category', 'id, name, description, iconName', '—', '→ Meditation'],
        ['UserProgress', 'totalSessions, totalMinutes, currentStreak, longestStreak', 'updateStreak(), addSession()', '→ User'],
        ['BreathingExercise', 'inhaleSeconds, holdSeconds, exhaleSeconds, cycles', 'getTotalDuration(), getPattern()', '—'],
    ],
    'Таблица 9 — Описание сущностей Domain Model'
)

add_h3('Бизнес-правила')
add_bullet('Пользователь не может иметь более одной записи настроения в сутки.')
add_bullet('Серия (streak) прерывается при пропуске хотя бы одного дня медитации.')
add_bullet('Медитация считается завершённой при прохождении более 80% времени.')
add_bullet('Оценка настроения принимает целочисленные значения от 1 до 10 включительно.')
add_bullet('Удаление медитации выполняется через soft delete (флаг active = false).')

add_h2('2.3. Архитектурное проектирование')

add_h3('2.3.1. Архитектурный паттерн PCMEF')
add_text(
    'Архитектурной основой системы служит паттерн PCMEF '
    '(Presentation-Control-Mediator-Entity-Foundation). Выбор паттерна '
    'обусловлен следующими требованиями: строгое разделение ответственности, '
    'тестируемость каждого слоя изолированно, направленность зависимостей '
    'только сверху вниз, адаптируемость под различные платформы (Android и JVM).'
)

add_simple_table(
    ['Слой', 'Ответственность', 'Backend', 'Android'],
    [
        ['P — Presentation', 'UI, отображение данных, обработка ввода', '—', 'Composable-функции, экраны'],
        ['C — Control', 'Координация, валидация, обработка запросов', 'REST-контроллеры (@RestController)', 'ViewModel + StateFlow'],
        ['M — Mediator', 'Бизнес-логика, транзакции, правила', 'Service-интерфейсы + реализации', 'Repository-интерфейсы + реализации'],
        ['E — Entity', 'Бизнес-объекты, состояние, методы предметной области', 'JPA-сущности с бизнес-методами', 'Data-классы Kotlin'],
        ['F — Foundation', 'Доступ к данным, инфраструктура', 'JPA-репозитории (Spring Data)', 'Room DAO + Retrofit API'],
    ],
    'Таблица 10 — Реализация слоёв PCMEF'
)

add_text(
    'Ключевое правило зависимостей: каждый слой может зависеть только от '
    'нижележащего. Обратные зависимости запрещены. Это гарантируется через '
    'использование интерфейсов на границах слоёв.'
)

add_h3('2.3.2. Интерфейсы между слоями (контракты)')
add_text('Интерфейс Control → Mediator (бэкенд):')
add_code('public interface MoodService {')
add_code('    MoodEntryDto save(Long userId, MoodEntryRequest request);')
add_code('    MoodEntryDto update(Long id, Long userId, MoodEntryRequest req);')
add_code('    List<MoodEntryDto> getHistory(Long userId, int days);')
add_code('    MoodEntryDto getToday(Long userId);')
add_code('    void delete(Long id, Long userId);')
add_code('    Double getAverage(Long userId, int days);')
add_code('}')

add_text('Интерфейс Mediator → Foundation (бэкенд):')
add_code('public interface MoodEntryRepository extends JpaRepository<MoodEntry, Long> {')
add_code('    List<MoodEntry> findByUserIdAndPeriod(Long userId, LocalDateTime from);')
add_code('    Optional<MoodEntry> findTodayEntry(Long userId, LocalDateTime start, LocalDateTime end);')
add_code('    Optional<Double> getAverageScore(Long userId, LocalDateTime from);')
add_code('}')

add_text('Интерфейс Control → Mediator (Android):')
add_code('interface IMoodRepository {')
add_code('    suspend fun saveMoodEntry(score: Int, note: String): MoodEntry')
add_code('    suspend fun getTodayEntry(): MoodEntry?')
add_code('    fun getHistory(days: Int): Flow<List<MoodEntry>>')
add_code('    suspend fun deleteEntry(id: Long)')
add_code('    suspend fun updateEntry(id: Long, score: Int, note: String): MoodEntry')
add_code('}')

add_h3('2.3.3. Архитектурные решения (ADR)')
add_simple_table(
    ['ADR', 'Решение', 'Обоснование'],
    [
        ['ADR-001', 'Паттерн PCMEF вместо Clean/Hexagonal', 'Требование методички, чёткая иерархия, адаптируется под обе платформы'],
        ['ADR-002', 'Java 17 + Spring Boot 3.2 (бэкенд)', 'LTS-версия, record-типы, sealed classes, широкая экосистема'],
        ['ADR-003', 'PostgreSQL 15 вместо MySQL', 'Лучшая поддержка JSON, оконные функции, ACID-совместимость'],
        ['ADR-004', 'JWT для аутентификации', 'Stateless, нет session storage, масштабируется горизонтально'],
        ['ADR-005', 'Jetpack Compose + Material Design 3', 'Декларативный UI, меньше бойлерплейта, современный toolkit'],
        ['ADR-006', 'Room SQLite для офлайн-кэша', 'Официальный Android ORM, интеграция с LiveData/Flow'],
        ['ADR-007', 'Retrofit 2 + OkHttp', 'Де-факто стандарт HTTP-клиента в Android'],
        ['ADR-008', 'Kotlin Coroutines + StateFlow', 'Структурированный асинхронный код, реактивные состояния'],
        ['ADR-009', 'Navigation Compose', 'Единый граф навигации, deep links, type-safe аргументы'],
    ],
    'Таблица 11 — Архитектурные решения (ADR)'
)

add_h2('2.4. Проектирование базы данных')

add_h3('2.4.1. Описание таблиц')
add_text(
    'База данных PostgreSQL 15 содержит 5 таблиц, нормализованных до третьей '
    'нормальной формы (3НФ). Все первичные ключи — суррогатные (BIGSERIAL), '
    'внешние ключи обеспечивают ссылочную целостность.'
)

add_simple_table(
    ['Таблица', 'Назначение', 'Ключевые поля'],
    [
        ['users', 'Пользователи системы', 'id (PK), email (UNIQUE), password, name, role'],
        ['categories', 'Категории медитаций', 'id (PK), name (UNIQUE), description, icon_name'],
        ['meditations', 'Каталог медитаций', 'id (PK), title, duration_minutes, difficulty_level, active, category_id (FK)'],
        ['meditation_sessions', 'Факты проведения медитаций', 'id (PK), started_at, completed, user_id (FK), meditation_id (FK)'],
        ['mood_entries', 'Записи дневника настроения', 'id (PK), score (1–10), note, recorded_at, user_id (FK)'],
    ],
    'Таблица 12 — Таблицы базы данных'
)

add_h3('2.4.2. Стратегия ORM')
add_text(
    'Для маппинга объектов на таблицы используется Spring Data JPA / Hibernate. '
    'Каждая таблица БД соответствует JPA-сущности в слое Entity. '
    'Репозитории в слое Foundation наследуют JpaRepository<Entity, Long> и '
    'предоставляют готовые CRUD-операции без написания SQL, а для сложных запросов '
    'используется @Query с JPQL.'
)
add_bullet('ddl-auto=update — Hibernate самостоятельно создаёт/обновляет схему при старте')
add_bullet('data.sql — инициализация тестовых данных (15 медитаций, 5 категорий)')
add_bullet('Lazy loading — связанные коллекции загружаются по требованию (FetchType.LAZY)')
add_bullet('Транзакции — @Transactional на уровне Service (Mediator-слой)')

add_h3('2.4.3. Фрагмент DDL-скрипта')
add_code('CREATE TABLE mood_entries (')
add_code('    id          BIGSERIAL PRIMARY KEY,')
add_code('    score       INTEGER NOT NULL CHECK (score >= 1 AND score <= 10),')
add_code('    note        TEXT,')
add_code('    recorded_at TIMESTAMP DEFAULT NOW(),')
add_code('    user_id     BIGINT NOT NULL REFERENCES users(id) ON DELETE CASCADE')
add_code(');')
add_code('CREATE INDEX idx_mood_user_date ON mood_entries(user_id, recorded_at DESC);')

add_h2('2.5. Детальное проектирование')

add_h3('2.5.1. Диаграмма последовательности: «Запись настроения»')
add_text(
    'Сценарий описывает процесс сохранения оценки настроения пользователем '
    'через мобильное приложение:'
)
add_bullet('Пользователь вводит оценку (1–10) и текстовую заметку на MoodDiaryScreen')
add_bullet('MoodViewModel вызывает IMoodRepository.saveMoodEntry(score, note)')
add_bullet('MoodRepositoryImpl сохраняет запись локально через Room DAO (оффлайн-гарантия)')
add_bullet('Retrofit-запрос POST /api/mood отправляется на сервер')
add_bullet('MoodController (Control) вызывает MoodService.save(userId, request)')
add_bullet('MoodServiceImpl (Mediator) создаёт MoodEntry-объект и сохраняет через Repository')
add_bullet('Ответ MoodEntryDto возвращается через цепочку назад до UI')
add_bullet('StateFlow обновляется, UI перерисовывается автоматически')

add_h3('2.5.2. Диаграмма последовательности: «Аутентификация»')
add_bullet('Пользователь вводит email и пароль на LoginScreen')
add_bullet('AuthViewModel → IAuthRepository.login(email, password)')
add_bullet('POST /api/auth/login → AuthController → AuthService.login()')
add_bullet('AuthServiceImpl проверяет BCrypt-хэш через PasswordEncoder')
add_bullet('JwtUtil генерирует accessToken (15 мин) и refreshToken (7 дней)')
add_bullet('Токены возвращаются в AuthResponse и сохраняются в TokenManager (DataStore)')
add_bullet('AuthInterceptor теперь будет добавлять Bearer-заголовок ко всем запросам')

add_h3('2.5.3. Паттерны GoF в проекте')
add_simple_table(
    ['Паттерн', 'Где применён', 'Назначение'],
    [
        ['Facade', 'Repository (Android)', 'Единая точка доступа к Room + Retrofit'],
        ['Template Method', 'Base ViewModel', 'Общий алгоритм загрузки с состояниями Loading/Success/Error'],
        ['Observer', 'StateFlow + Compose', 'Реактивная связь ViewModel и UI'],
        ['Strategy', 'Difficulty enum в Meditation', 'Алгоритмы выбора контента по сложности'],
        ['Singleton', 'Room Database, Retrofit', 'Единственный экземпляр через DI контейнер'],
    ],
    'Таблица 13 — Применённые паттерны GoF'
)
page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# ГЛАВА 3. РЕАЛИЗАЦИОННАЯ ЧАСТЬ
# ═══════════════════════════════════════════════════════════════════════════════
add_h1('3. РЕАЛИЗАЦИОННАЯ ЧАСТЬ')

add_h2('3.1. Структура проекта')
add_text(
    'Репозиторий организован как монорепозиторий с чёткой структурой директорий:'
)
add_code('MindFlow/')
add_code('├── android-app/          # Kotlin/Compose Android-приложение')
add_code('│   └── app/src/main/java/ru/mindflow/app/')
add_code('│       ├── presentation/  # P-слой: Compose экраны')
add_code('│       ├── control/       # C-слой: ViewModels')
add_code('│       ├── mediator/      # M-слой: Repository-интерфейсы')
add_code('│       ├── entity/        # E-слой: Domain модели')
add_code('│       └── foundation/    # F-слой: Room + Retrofit')
add_code('├── backend/              # Java Spring Boot сервер')
add_code('│   └── src/main/java/ru/mindflow/backend/')
add_code('│       ├── control/       # REST Controllers')
add_code('│       ├── mediator/      # Services')
add_code('│       ├── entity/        # JPA Entities')
add_code('│       ├── foundation/    # JPA Repositories')
add_code('│       ├── dto/           # Data Transfer Objects')
add_code('│       ├── security/      # JWT + Spring Security')
add_code('│       └── config/        # SecurityConfig, JwtConfig')
add_code('├── docs/                 # Документация (13 разделов)')
add_code('├── docker-compose.yml    # Оркестрация контейнеров')
add_code('└── README.md')

add_h2('3.2. Реализация серверной части (Backend)')

add_h3('3.2.1. Слой Entity (E)')
add_text(
    'JPA-сущности расположены в пакете ru.mindflow.backend.entity. '
    'Ключевое требование: сущности не должны быть анемичными — они содержат '
    'бизнес-методы, отражающие логику предметной области.'
)
add_text('Пример: бизнес-метод класса MoodEntry:')
add_code('public String getMoodLabel() {')
add_code('    return switch (score) {')
add_code('        case 9, 10 -> "Отлично";')
add_code('        case 7, 8  -> "Хорошо";')
add_code('        case 5, 6  -> "Нормально";')
add_code('        case 3, 4  -> "Плохо";')
add_code('        default    -> "Очень плохо";')
add_code('    };')
add_code('}')
add_text(
    'Все сущности аннотированы @Entity, @Table и используют Lombok для '
    'генерации конструкторов, геттеров и сеттеров. Связи между сущностями '
    'реализованы через @ManyToOne / @OneToMany с FetchType.LAZY.'
)

add_h3('3.2.2. Слой Foundation (F)')
add_text(
    'Репозитории расположены в пакете ru.mindflow.backend.foundation. '
    'Каждый репозиторий наследует JpaRepository<Entity, Long>, что '
    'обеспечивает готовые CRUD-операции. Для сложных запросов используется '
    'аннотация @Query с JPQL-синтаксисом:'
)
add_code('@Query("""')
add_code('    SELECT AVG(m.score) FROM MoodEntry m')
add_code('    WHERE m.user.id = :userId')
add_code('    AND m.recordedAt >= :from')
add_code('""")')
add_code('Optional<Double> getAverageScore(')
add_code('    @Param("userId") Long userId,')
add_code('    @Param("from") LocalDateTime from);')

add_h3('3.2.3. Слой Mediator — Сервисы (M)')
add_text(
    'Бизнес-логика инкапсулирована в сервисах, реализующих интерфейсы. '
    'Разделение интерфейс/реализация обеспечивает возможность тестирования '
    'через мокирование и соответствует принципу инверсии зависимостей (DIP).'
)

add_simple_table(
    ['Интерфейс', 'Реализация', 'Методы'],
    [
        ['AuthService', 'AuthServiceImpl', 'register(), login(), refresh(), buildAuthResponse()'],
        ['MeditationService', 'MeditationServiceImpl', 'getAll(), getById(), getByCategory(), search()'],
        ['MoodService', 'MoodServiceImpl', 'save(), update(), getHistory(), getToday(), getAverage(), delete()'],
        ['CategoryService', 'CategoryServiceImpl', 'findAll(), findById()'],
    ],
    'Таблица 14 — Сервисы слоя Mediator'
)

add_text(
    'Все методы, изменяющие состояние БД, аннотированы @Transactional. '
    'Методы только для чтения используют @Transactional(readOnly = true) '
    'для оптимизации производительности.'
)

add_h3('3.2.4. Слой Control — REST-контроллеры (C)')
add_text(
    'Контроллеры обрабатывают HTTP-запросы, выполняют валидацию входных данных '
    '(@Valid) и делегируют бизнес-логику сервисам. Контроллеры не содержат '
    'бизнес-логики — это нарушило бы принцип разделения ответственности PCMEF.'
)

add_simple_table(
    ['Контроллер', 'Путь', 'Методы | Описание'],
    [
        ['AuthController', '/api/auth', 'POST register | Регистрация\nPOST login | Аутентификация\nPOST refresh | Обновление токена'],
        ['MeditationController', '/api/meditations', 'GET / | Каталог (фильтр/поиск)\nGET /{id} | Детали медитации'],
        ['CategoryController', '/api/categories', 'GET / | Список категорий\nGET /{id} | Категория по ID'],
        ['MoodController', '/api/mood', 'POST / | Создать запись\nGET / | История\nPUT /{id} | Обновить\nGET /today | За сегодня\nGET /average | Средняя\nDELETE /{id} | Удалить'],
    ],
    'Таблица 15 — REST-контроллеры системы'
)

add_h2('3.3. Реализация Android-клиента')

add_h3('3.3.1. Слой Foundation — локальное и сетевое хранилище (F)')
add_text(
    'Фундаментальный слой Android включает два источника данных: '
    'локальный (Room/SQLite) и удалённый (Retrofit/REST API).'
)
add_text('Room Database:')
add_code('@Database(entities = [MeditationEntity::class, MoodEntryEntity::class,')
add_code('                       GardenEntity::class], version = 1)')
add_code('abstract class MindFlowDatabase : RoomDatabase() {')
add_code('    abstract fun meditationDao(): MeditationDao')
add_code('    abstract fun moodEntryDao(): MoodEntryDao')
add_code('    abstract fun gardenDao(): GardenDao')
add_code('}')

add_text('Retrofit API-интерфейс:')
add_code('interface MindFlowApi {')
add_code('    @GET("meditations")')
add_code('    suspend fun getMeditations(@Query("categoryId") id: Long? = null,')
add_code('                               @Query("search") q: String? = null): List<MeditationDto>')
add_code('    @POST("mood")')
add_code('    suspend fun postMood(@Body request: MoodEntryRequest): MoodEntryDto')
add_code('}')

add_text(
    'TokenManager хранит JWT-токены в зашифрованном DataStore Preferences, '
    'обеспечивая безопасное персистентное хранение на устройстве.'
)

add_h3('3.3.2. Слой Mediator — репозитории (M)')
add_text(
    'Репозитории реализуют стратегию offline-first: данные всегда читаются '
    'из локального кэша, а сетевые запросы синхронизируют актуальное состояние.'
)
add_code('class MeditationRepositoryImpl(')
add_code('    private val api: MindFlowApi,')
add_code('    private val dao: MeditationDao')
add_code(') : IMeditationRepository {')
add_code('    override fun getAllMeditations(): Flow<List<Meditation>> = flow {')
add_code('        emit(dao.getAll().map { it.toDomain() })  // сначала кэш')
add_code('        val remote = api.getMeditations()')
add_code('        dao.upsertAll(remote.map { it.toEntity() }) // обновить кэш')
add_code('        emit(dao.getAll().map { it.toDomain() })   // обновлённые данные')
add_code('    }')
add_code('}')

add_h3('3.3.3. Слой Control — ViewModels (C)')
add_text(
    'ViewModels управляют состоянием экранов через StateFlow. '
    'Все операции выполняются в Coroutine Scope, что обеспечивает '
    'корректную работу с жизненным циклом Android-компонентов.'
)
add_code('class MoodViewModel(private val repo: IMoodRepository) : ViewModel() {')
add_code('    private val _state = MutableStateFlow<MoodUiState>(MoodUiState.Idle)')
add_code('    val state: StateFlow<MoodUiState> = _state.asStateFlow()')
add_code('')
add_code('    fun saveMoodEntry(score: Int, note: String) {')
add_code('        viewModelScope.launch {')
add_code('            _state.value = MoodUiState.Loading')
add_code('            _state.value = try {')
add_code('                val entry = repo.saveMoodEntry(score, note)')
add_code('                MoodUiState.Success(entry)')
add_code('            } catch (e: Exception) {')
add_code('                MoodUiState.Error(e.message ?: "Ошибка")')
add_code('            }')
add_code('        }')
add_code('    }')
add_code('}')

add_h3('3.3.4. Слой Presentation — экраны Compose (P)')
add_text('Приложение содержит 12 экранов, реализованных с Jetpack Compose:')
add_simple_table(
    ['Экран', 'Маршрут', 'Функционал', 'ViewModel'],
    [
        ['LoginScreen', '/login', 'Email + пароль, ссылка на регистрацию', 'AuthViewModel'],
        ['RegisterScreen', '/register', 'Форма регистрации (3 поля)', 'AuthViewModel'],
        ['WelcomeScreen', '/welcome', 'Онбординг-экран перед входом', '—'],
        ['HomeScreen', '/home', 'Приветствие, Quick Actions, статистика', 'HomeViewModel'],
        ['MeditationListScreen', '/meditations', 'Список медитаций, поиск, фильтр', 'MeditationViewModel'],
        ['MeditationDetailScreen', '/meditation/{id}', 'Детали медитации, кнопка «Начать»', 'MeditationViewModel'],
        ['MeditationSessionScreen', '/session/{id}', 'Таймер, прогресс-бар, «Завершить»', 'MeditationViewModel'],
        ['MoodDiaryScreen', '/mood', 'Слайдер 1–10, заметка, «Сохранить»', 'MoodViewModel'],
        ['AnalyticsScreen', '/analytics', 'График настроения за период', 'MoodViewModel'],
        ['ProfileScreen', '/profile', 'Информация, прогресс, «Выйти»', 'ProfileViewModel'],
        ['GardenScreen', '/garden', 'Геймификация: цветы, уровни', 'GardenViewModel'],
        ['CourseScreen', '/courses', 'Курсы и программы практик', '—'],
    ],
    'Таблица 16 — Экраны Android-приложения'
)

add_h3('3.3.5. Навигация')
add_text(
    'Навигация реализована через Navigation Compose с единым NavGraph. '
    'Нижняя панель навигации (BottomNavigationBar) обеспечивает переход '
    'между четырьмя основными разделами: Home, Медитации, Настроение, Профиль.'
)

add_h2('3.4. Рефакторинг и паттерны качества')

add_h3('3.4.1. Паттерн Data Mapper')
add_text(
    'Паттерн Data Mapper отделяет бизнес-логику доменных объектов от '
    'механизма доступа к данным. В проекте реализован через extension-функции '
    'toEntity() / toDomain() / toDto(), разделяющие три представления данных:'
)
add_bullet('Domain-объект (Entity-слой): чистая бизнес-логика, без зависимостей от инфраструктуры')
add_bullet('Room Entity (Foundation-слой): оптимизирован для SQLite-хранения')
add_bullet('DTO (Control-слой): сериализуемое представление для REST API')
add_code('// Android: Data Mapper через extension-функции')
add_code('fun MeditationEntity.toDomain(): Meditation = Meditation(')
add_code('    id = this.id, title = this.title,')
add_code('    durationMinutes = this.durationMinutes, ...')
add_code(')')
add_code('fun MeditationDto.toEntity(): MeditationEntity = MeditationEntity(')
add_code('    id = this.id, title = this.title, ...')
add_code(')')

add_h3('3.4.2. Паттерн Identity Map')
add_text(
    'Identity Map гарантирует уникальность объектов в рамках сессии — '
    'один и тот же объект по одному и тому же идентификатору должен '
    'существовать в единственном экземпляре. В проекте реализован двумя способами:'
)
add_bullet('Backend: Hibernate First-Level Cache — в рамках одной @Transactional сессии сущность по ID загружается из БД только один раз')
add_bullet('Android (Room DAO): метод upsertAll() с @Upsert аннотацией — при получении медитации с уже существующим ID обновляет запись, а не дублирует её')
add_code('// Android: Identity Map через Room @Upsert')
add_code('@Upsert')
add_code('suspend fun upsertAll(meditations: List<MeditationEntity>)')

add_h3('3.4.3. Результаты статического анализа')
add_simple_table(
    ['Компонент', 'Инструмент', 'Выявленные проблемы', 'Исправлено'],
    [
        ['Backend', 'Checkstyle + IntelliJ', '3 неиспользуемых импорта, 2 magic number', '5 из 5'],
        ['Android', 'Android Lint', '1 missing content description, 2 hardcoded strings', '3 из 3'],
        ['Общее', 'SonarLint IDE plugin', '0 Critical, 2 Major (устранены)', '2 из 2'],
    ],
    'Таблица 17 — Результаты статического анализа'
)

add_h2('3.5. Безопасность системы')
add_text(
    'Система реализует многоуровневую защиту данных пользователей:'
)
add_simple_table(
    ['Уровень', 'Механизм', 'Реализация'],
    [
        ['Хранение паролей', 'BCrypt (strength=12)', 'Spring Security PasswordEncoder'],
        ['Аутентификация', 'JWT Access Token (15 мин)', 'JJWT 0.12.3, HMAC-SHA256'],
        ['Refresh сессии', 'JWT Refresh Token (7 дней)', 'POST /api/auth/refresh'],
        ['Защита эндпоинтов', 'Spring Security Filter Chain', 'JwtAuthFilter (OncePerRequestFilter)'],
        ['Хранение токенов', 'Encrypted DataStore', 'Android DataStore Preferences'],
        ['Транспорт токена', 'Authorization: Bearer <token>', 'OkHttp Interceptor (AuthInterceptor)'],
        ['Ролевая модель', 'ROLE_USER / ROLE_ADMIN', '@PreAuthorize в контроллерах'],
    ],
    'Таблица 18 — Механизмы безопасности'
)

add_h2('3.6. REST API')
add_text(
    'REST API реализован в соответствии с принципами RESTful-архитектуры. '
    'Документация автоматически генерируется через библиотеку springdoc-openapi '
    'и доступна в виде интерактивного Swagger UI.'
)
add_simple_table(
    ['Метод', 'Эндпоинт', 'Описание', 'Доступ'],
    [
        ['POST', '/api/auth/register', 'Регистрация нового пользователя', 'Публичный'],
        ['POST', '/api/auth/login', 'Аутентификация, получение токенов', 'Публичный'],
        ['POST', '/api/auth/refresh', 'Обновление access-token', 'Публичный'],
        ['GET', '/api/categories', 'Список всех категорий медитаций', 'USER, ADMIN'],
        ['GET', '/api/categories/{id}', 'Категория по идентификатору', 'USER, ADMIN'],
        ['GET', '/api/meditations', 'Каталог медитаций (с фильтром)', 'USER, ADMIN'],
        ['GET', '/api/meditations/{id}', 'Детали медитации', 'USER, ADMIN'],
        ['POST', '/api/mood', 'Создать запись настроения', 'USER'],
        ['GET', '/api/mood', 'История настроения (с параметром days)', 'USER'],
        ['PUT', '/api/mood/{id}', 'Обновить запись настроения', 'USER'],
        ['GET', '/api/mood/today', 'Запись настроения за сегодня', 'USER'],
        ['GET', '/api/mood/average', 'Средняя оценка за период', 'USER'],
        ['DELETE', '/api/mood/{id}', 'Удалить запись настроения', 'USER'],
    ],
    'Таблица 19 — REST API эндпоинты'
)
add_text(
    'Все защищённые эндпоинты требуют заголовок Authorization: Bearer <access_token>. '
    'Стандартные HTTP-коды: 200 OK, 201 Created, 204 No Content, 400 Bad Request, '
    '401 Unauthorized, 403 Forbidden, 404 Not Found, 409 Conflict, 500 Server Error.'
)
page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# ГЛАВА 4. ТЕСТИРОВАНИЕ
# ═══════════════════════════════════════════════════════════════════════════════
add_h1('4. ТЕСТИРОВАНИЕ И ОБЕСПЕЧЕНИЕ КАЧЕСТВА')

add_h2('4.1. Стратегия тестирования')
add_text(
    'Стратегия тестирования включает три уровня: модульное тестирование '
    'отдельных классов, интеграционное тестирование слоёв и ручное '
    'системное тестирование по тест-кейсам. Целевое покрытие — более 40% '
    'строк кода на каждой платформе (требование методички).'
)
add_simple_table(
    ['Уровень', 'Инструмент', 'Объект тестирования', 'Результат'],
    [
        ['Модульные (Backend)', 'JUnit 5 + Mockito + MockMvc', 'Service, Controller, Entity', '36 тестов, 58,1% LINE'],
        ['Модульные (Android)', 'JUnit 4 + kotlinx-coroutines-test', 'ViewModel, Repository, Entity', '253 теста, 43,6% LINE'],
        ['Ручное системное', 'Тест-кейсы TC-01…TC-12', 'E2E пользовательские сценарии', '12 из 12 пройдено'],
    ],
    'Таблица 20 — Уровни тестирования'
)

add_h2('4.2. Модульное тестирование бэкенда')
add_text(
    'Тестирование серверной части выполнено с использованием JUnit 5, Mockito '
    'и MockMvc (Spring Boot Test). База данных в тестах — H2 in-memory '
    '(режим совместимости с PostgreSQL).'
)
add_simple_table(
    ['Класс', 'Тестов', 'Что проверяется'],
    [
        ['MoodServiceImplTest', '12', 'save(), update(), delete(), getHistory(), getToday(), getAverage()'],
        ['MeditationServiceImplTest', '10', 'getAll(), getById(), search() по запросу и категории'],
        ['AuthServiceImplTest', '7', 'register() с дубликатом, login() с неверным паролем, refresh()'],
        ['MoodEntryEntityTest', '6', 'getMoodLabel() для всех диапазонов баллов'],
        ['MoodControllerTest', '1+', 'HTTP статусы, сериализация JSON'],
    ],
    'Таблица 21 — Тесты серверной части'
)

add_text('Отчёт JaCoCo (Backend):')
add_simple_table(
    ['Метрика', 'Покрыто', 'Всего', 'Процент'],
    [
        ['LINE', '125', '215', '58,1%'],
        ['METHOD', '34', '60', '56,7%'],
        ['CLASS', '12', '18', '66,7%'],
        ['INSTRUCTION', '862', '1 372', '62,8%'],
    ],
    'Таблица 22 — Покрытие тестами (Backend)'
)

add_h2('4.3. Модульное тестирование Android')
add_text(
    '253 модульных теста охватывают все основные компоненты Android-приложения. '
    'Тесты выполняются на JVM без запуска эмулятора (unit tests), что '
    'обеспечивает быстрое выполнение в CI-пайплайне.'
)
add_simple_table(
    ['Пакет', 'Тестов', 'Что тестируется'],
    [
        ['entity', '~30', 'Бизнес-методы Meditation, MoodEntry, User'],
        ['garden', '34', 'GardenProgress: уровни, цветы, декорации'],
        ['control', '~80', 'AuthViewModel, MoodViewModel, MeditationViewModel, HomeViewModel'],
        ['mediator', '~60', 'AuthRepositoryImpl, MoodRepositoryImpl, MeditationRepositoryImpl'],
        ['foundation.local.entity', '~30', 'Room Entity классы и их маппинг'],
        ['foundation.remote.dto', '~20', 'DTO сериализация/десериализация'],
        ['presentation.navigation', '~10', 'Screen маршруты и аргументы'],
        ['Прочие', '~19', 'AppContainer, утилитарные классы'],
    ],
    'Таблица 23 — Тесты Android-приложения'
)

add_text('Отчёт JaCoCo (Android):')
add_simple_table(
    ['Метрика', 'Покрыто', 'Всего', 'Процент'],
    [
        ['LINE', '331', '760', '43,6%'],
        ['METHOD', '177', '243', '72,8%'],
        ['CLASS', '29', '35', '82,9%'],
        ['COMPLEXITY', '211', '334', '63,2%'],
    ],
    'Таблица 24 — Покрытие тестами (Android)'
)

add_h2('4.4. Ручное тестирование')
add_simple_table(
    ['TC', 'Тест-кейс', 'Ожидаемый результат', 'Статус'],
    [
        ['TC-01', 'Регистрация с валидными данными', '200 OK, токены в ответе', '✅ PASS'],
        ['TC-02', 'Регистрация с существующим email', '409 Conflict', '✅ PASS'],
        ['TC-03', 'Вход с неверным паролем', '401 Unauthorized', '✅ PASS'],
        ['TC-04', 'Обращение без токена к /api/mood', '401 Unauthorized', '✅ PASS'],
        ['TC-05', 'Загрузка каталога медитаций', 'Список из 15 медитаций', '✅ PASS'],
        ['TC-06', 'Фильтрация медитаций по категории', 'Только медитации данной категории', '✅ PASS'],
        ['TC-07', 'Поиск медитации по тексту', 'Релевантные результаты', '✅ PASS'],
        ['TC-08', 'Запись настроения (score=7)', '200 OK, label="Хорошо"', '✅ PASS'],
        ['TC-09', 'Обновление записи настроения', '200 OK, обновлённые данные', '✅ PASS'],
        ['TC-10', 'Запрос среднего за 30 дней', 'Double значение, ≥1.0 и ≤10.0', '✅ PASS'],
        ['TC-11', 'Работа в офлайн-режиме', 'Медитации доступны из Room кэша', '✅ PASS'],
        ['TC-12', 'Refresh token при истёкшем access', 'Новый access token без повторного входа', '✅ PASS'],
    ],
    'Таблица 25 — Результаты ручного тестирования'
)
page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# ГЛАВА 5. РАЗВЁРТЫВАНИЕ
# ═══════════════════════════════════════════════════════════════════════════════
add_h1('5. РАЗВЁРТЫВАНИЕ И ЭКСПЛУАТАЦИЯ')

add_h2('5.1. Контейнеризация (Docker)')
add_text(
    'Серверная часть системы контейнеризована с использованием Docker. '
    'Файл docker-compose.yml описывает два сервиса: база данных PostgreSQL '
    'и Spring Boot-сервер. Сервис бэкенда зависит от готовности БД '
    '(condition: service_healthy).'
)
add_code('services:')
add_code('  postgres:')
add_code('    image: postgres:15-alpine')
add_code('    environment:')
add_code('      POSTGRES_DB: mindflow_db')
add_code('      POSTGRES_USER: mindflow_user')
add_code('      POSTGRES_PASSWORD: mindflow_pass')
add_code('    ports: ["5433:5432"]')
add_code('')
add_code('  backend:')
add_code('    build: ./backend')
add_code('    ports: ["8081:8081"]')
add_code('    depends_on:')
add_code('      postgres: { condition: service_healthy }')

add_h2('5.2. Инструкция по развёртыванию')
add_text('Способ 1 — через Docker (рекомендуется):')
add_bullet('Установить Docker Desktop')
add_bullet('Выполнить: docker-compose up -d')
add_bullet('Swagger UI доступен: http://localhost:8081/swagger-ui.html')

add_text('Способ 2 — ручная установка:')
add_bullet('Создать БД PostgreSQL: CREATE DATABASE mindflow_db;')
add_bullet('Сборка: cd backend && ./mvnw clean package -DskipTests')
add_bullet('Запуск JAR с переменными окружения (см. docs/11-user-guide/admin-guide.md)')

add_text('Android APK:')
add_bullet('Android Studio: Build → Generate Signed APK')
add_bullet('Командная строка: ./gradlew assembleDebug')
add_bullet('Файл: android-app/app/build/outputs/apk/debug/app-debug.apk')

add_h2('5.3. Требования к окружению')
add_simple_table(
    ['Компонент', 'Минимальная версия', 'Рекомендуемая'],
    [
        ['Java JDK', '17', '21'],
        ['Maven', '3.8', '3.9'],
        ['PostgreSQL', '14', '15'],
        ['Docker', '20.10', '24.x'],
        ['Android Studio', 'Hedgehog 2023.1.1', 'Meerkat 2024.3.1'],
        ['Android SDK', 'API 26 (Android 8.0)', 'API 35 (Android 15)'],
    ],
    'Таблица 26 — Требования к окружению'
)
page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# ГЛАВА 6. УПРАВЛЕНИЕ ПРОЕКТОМ
# ═══════════════════════════════════════════════════════════════════════════════
add_h1('6. УПРАВЛЕНИЕ ПРОЕКТОМ')

add_h2('6.1. WBS — иерархическая структура работ')
add_text(
    'Иерархическая структура работ (Work Breakdown Structure) описывает '
    'декомпозицию проекта на 10 пакетов работ верхнего уровня:'
)
add_simple_table(
    ['Код', 'Пакет работ', 'Трудоёмкость (ч)', 'Доля'],
    [
        ['1.1', 'Управление проектом', '8', '5%'],
        ['1.2', 'Аналитика и требования', '14', '9%'],
        ['1.3', 'Архитектурное проектирование', '12', '8%'],
        ['1.4', 'Проектирование БД', '8', '5%'],
        ['1.5', 'Серверная часть (Backend)', '30', '19%'],
        ['1.6', 'Android-приложение', '40', '26%'],
        ['1.7', 'Тестирование', '14', '9%'],
        ['1.8', 'Рефакторинг и качество', '8', '5%'],
        ['1.9', 'Развёртывание', '6', '4%'],
        ['1.10', 'Документация', '16', '10%'],
        ['**ИТОГО**', '', '**156**', '**100%**'],
    ],
    'Таблица 27 — WBS (верхний уровень)'
)

add_h2('6.2. Диаграмма Ганта')
add_text(
    'Разработка велась по итеративному плану в соответствии с 18-недельным '
    'графиком методички. Фактические сроки были несколько сжаты (~ 6 недель '
    'активной разработки, 01.05–09.06.2026) за счёт интенсивного темпа работы.'
)
add_simple_table(
    ['Этап', 'Неделя', 'Продолжительность', 'Ключевые артефакты'],
    [
        ['Этап 0: Инициация', '1–2', '2 нед.', 'Паспорт, IDEF0, BUC, SWOT, CJM, глоссарий'],
        ['Этап 1: Требования', '3–4', '2 нед.', 'Use Case, Domain Model, трассировка'],
        ['Этап 2: Архитектура', '5–6', '2 нед.', 'PCMEF, ADR, интерфейсы'],
        ['Этап 3: БД', '7–8', '2 нед.', 'ER, DDL, ORM-стратегия'],
        ['Этап 4: Детал. проектирование', '9–10', '2 нед.', 'Sequence, Class diagrams'],
        ['Этап 5: Реализация ядра', '11–12', '2 нед.', 'Backend + Android core'],
        ['Этап 6: Рефакторинг', '13–14', '2 нед.', 'Тесты, покрытие >40%, паттерны'],
        ['Этап 7: Интерфейс', '15–16', '2 нед.', '12 экранов, Material Design 3'],
        ['Этап 8: Завершение', '17–18', '2 нед.', 'Документация, Docker, презентация'],
    ],
    'Таблица 28 — Диаграмма Ганта (сводная)'
)

add_h2('6.3. Оценка трудозатрат — COCOMO I Basic')
add_text(
    'Для оценки трудозатрат применялась модель COCOMO I в режиме Embedded '
    '(встроенные системы), так как проект имеет нетривиальные архитектурные '
    'ограничения и требования к безопасности.'
)
add_simple_table(
    ['Параметр', 'Значение'],
    [
        ['Объём кода (KDSI)', '9,0 тыс. строк'],
        ['Режим COCOMO', 'Embedded (a=3.6, b=1.20)'],
        ['Расчётные трудозатраты E', '3,6 × 9,0^1,20 ≈ 50,3 ч/м'],
        ['Фактические трудозатраты', '~156 часов'],
        ['EAF (фактор корректировки)', '≈ 1,51 (студент, сложный стек)'],
        ['Скоррект. трудозатраты', '≈ 76 нормо-часов/мес'],
        ['Условная стоимость', '≈ 93 600 руб. (156 ч × 600 руб./ч)'],
    ],
    'Таблица 29 — Расчёт по COCOMO'
)
add_text(
    'Фактические 156 часов соответствуют скорректированной оценке COCOMO '
    'для одного студента-разработчика с учётом коэффициентов сложности '
    'и опыта разработчика.'
)
page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# ЗАКЛЮЧЕНИЕ
# ═══════════════════════════════════════════════════════════════════════════════
add_h1('ЗАКЛЮЧЕНИЕ')

add_text(
    'В ходе выполнения курсового проекта было разработано мобильное приложение '
    '«MindFlow» для поддержки ментального здоровья. Проект реализован по '
    'траектории В (мобильная разработка) и представляет собой полноценную '
    'клиент-серверную систему.'
)
add_text('Достигнуты следующие результаты:')
add_bullet('Разработан Android-клиент на Kotlin/Jetpack Compose с 12 экранами и поддержкой офлайн-режима через Room')
add_bullet('Реализована серверная часть на Java 17 / Spring Boot 3.x с 15 REST API эндпоинтами и документацией OpenAPI/Swagger')
add_bullet('Спроектирована и реализована архитектура на паттерне PCMEF с чёткими интерфейсами между слоями на обеих платформах')
add_bullet('Обеспечена безопасность: JWT-аутентификация, BCrypt-хеширование паролей, ролевая модель (ROLE_USER / ROLE_ADMIN)')
add_bullet('Написано 289 модульных тестов (36 backend + 253 Android) с покрытием 58,1% и 43,6% соответственно')
add_bullet('Реализованы паттерны качества Data Mapper и Identity Map')
add_bullet('Подготовлен полный комплект документации в 13 разделах: от паспорта проекта до пояснительной записки')
add_bullet('Система развёртывается через Docker Compose одной командой')

add_text(
    'В процессе разработки были получены практические навыки проектирования '
    'клиент-серверных мобильных приложений, применения архитектурного паттерна '
    'PCMEF, работы с JWT-аутентификацией, написания модульных тестов и '
    'документирования программного обеспечения.'
)
add_text(
    'Перспективы развития проекта: добавление push-уведомлений (Firebase FCM), '
    'биометрической аутентификации, расширенной аналитики с рекомендательной '
    'системой, публикация в Google Play Store и реализация Premium-подписки.'
)
page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# СПИСОК ИСТОЧНИКОВ
# ═══════════════════════════════════════════════════════════════════════════════
add_h1('СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ')

sources = [
    'Мацяшек, Л. Практическая программная инженерия на основе учебного примера / Л. Мацяшек. — М.: Вильямс, 2021. — 784 с.',
    'Мартин, Р. Чистая архитектура. Искусство разработки программного обеспечения / Р. Мартин. — СПб.: Питер, 2018. — 352 с.',
    'Фаулер, М. Шаблоны корпоративных приложений / М. Фаулер. — М.: Вильямс, 2016. — 544 с.',
    'Соммервилл, И. Инженерия программного обеспечения / И. Соммервилл. — 10-е изд. — М.: Вильямс, 2018. — 1008 с.',
    'Скотт, Д. Kotlin в действии / Д. Скотт, С. Эрмэн. — СПб.: ДМК Пресс, 2021. — 360 с.',
    'Spring Boot Reference Documentation. — [Электронный ресурс]. — Режим доступа: https://docs.spring.io/spring-boot/docs/current/reference/html/, свободный. — (дата обращения: 20.05.2026).',
    'Jetpack Compose Documentation. — [Электронный ресурс]. — Режим доступа: https://developer.android.com/jetpack/compose, свободный. — (дата обращения: 20.05.2026).',
    'OpenAPI Specification 3.0. — [Электронный ресурс]. — Режим доступа: https://swagger.io/specification/, свободный. — (дата обращения: 15.05.2026).',
    'OWASP Mobile Security Testing Guide. — [Электронный ресурс]. — Режим доступа: https://owasp.org/www-project-mobile-security-testing-guide/, свободный. — (дата обращения: 18.05.2026).',
    'Кабан, О.С. Методические указания по выполнению курсового проекта по дисциплине «Программная инженерия» / О.С. Кабан. — Ставрополь: СКФУ, 2026. — 48 с.',
]

for i, src in enumerate(sources, 1):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f'{i}. {src}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# ПРИЛОЖЕНИЯ
# ═══════════════════════════════════════════════════════════════════════════════
add_h1('ПРИЛОЖЕНИЯ')

add_h2('Приложение А. Структура каталогов репозитория')
add_code('MindFlow/')
add_code('├── android-app/')
add_code('│   ├── app/src/main/java/ru/mindflow/app/')
add_code('│   │   ├── presentation/screen/   (12 Compose экранов)')
add_code('│   │   ├── control/               (6 ViewModel)')
add_code('│   │   ├── mediator/              (3 Repository интерфейса + реализации)')
add_code('│   │   ├── entity/                (4 Domain класса)')
add_code('│   │   └── foundation/            (Room + Retrofit + TokenManager)')
add_code('│   └── app/src/test/              (253 JUnit тестов)')
add_code('├── backend/src/main/java/ru/mindflow/backend/')
add_code('│   ├── control/                   (4 REST контроллера)')
add_code('│   ├── mediator/                  (4 Service интерфейса + реализации)')
add_code('│   ├── entity/                    (5 JPA сущностей)')
add_code('│   ├── foundation/                (5 JPA репозиториев)')
add_code('│   ├── dto/                       (6 DTO классов)')
add_code('│   └── security/                  (JWT + Spring Security)')
add_code('├── docs/')
add_code('│   ├── 00-project-charter/        (паспорт, IDEF0, SWOT, CJM, ROI)')
add_code('│   ├── 01-requirements/           (Use Case, Domain Model, трассировка)')
add_code('│   ├── 02-architecture/           (PCMEF, Arc42, ADR, интерфейсы)')
add_code('│   ├── 03-database/               (ER, DDL, ORM-стратегия)')
add_code('│   ├── 04-detailed-design/        (Sequence, Class диаграммы)')
add_code('│   ├── 05-implementation/         (отчёт о реализации)')
add_code('│   ├── 06-testing/                (тест-план, JaCoCo отчёты)')
add_code('│   ├── 07-refactoring/            (анализ, паттерны, отчёт)')
add_code('│   ├── 08-ui/                     (описание UI, скриншоты)')
add_code('│   ├── 09-api/                    (OpenAPI спецификация)')
add_code('│   ├── 10-deployment/             (Docker, инструкция)')
add_code('│   ├── 11-user-guide/             (user + admin руководства)')
add_code('│   └── 12-final-report/           (ТЗ, WBS, Ганта, COCOMO, ПЗ)')
add_code('├── docker-compose.yml')
add_code('└── README.md')

add_h2('Приложение Б. Метрики проекта')
add_simple_table(
    ['Метрика', 'Значение'],
    [
        ['Коммиты в репозитории', '61'],
        ['Период разработки', '01.05.2026 – 09.06.2026'],
        ['Строк кода (SLOC)', '~9 000'],
        ['Экранов Android-приложения', '12'],
        ['REST API эндпоинтов', '15 (13 основных + 2 категории)'],
        ['Тестов (Backend)', '36'],
        ['Тестов (Android)', '253'],
        ['Покрытие Backend (LINE)', '58,1%'],
        ['Покрытие Android (LINE)', '43,6%'],
        ['Разделов документации', '13'],
        ['ADR (архитектурных решений)', '9'],
        ['Таблиц в БД', '5'],
        ['Медитаций в тестовых данных', '15'],
        ['Категорий медитаций', '5'],
    ],
    'Таблица Б.1 — Итоговые метрики проекта'
)

# ═══════════════════════════════════════════════════════════════════════════════
# СОХРАНЕНИЕ
# ═══════════════════════════════════════════════════════════════════════════════
output_path = 'Пояснительная_записка_MindFlow.docx'
doc.save(output_path)
print('OK: ' + output_path)